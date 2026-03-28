#!/usr/bin/env python3
"""
EA Assistant — Word Document Generator
Produces EA-structured .docx files from artifact content JSON.

Usage:
    python3 generate-docx.py --type vision --title "Architecture Vision" \
        --org "Acme Corp" --architect "Jane Smith" \
        --output "architecture-vision.docx" --content '{"sections": [...]}'

    python3 generate-docx.py --type vision --engagement-dir EA-projects/acme-2024 \
        --output "architecture-vision.docx"

Requirements: pip3 install python-docx
"""

import argparse
import json
import os
import sys
from datetime import date, datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)

try:
    from lxml import etree as _etree
    _HAS_LXML = True
except ImportError:
    _HAS_LXML = False

DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE  = RGBColor(0x2E, 0x74, 0xB5)
ORANGE    = RGBColor(0xC5, 0x5A, 0x11)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

ARTIFACT_SECTIONS = {
    "vision": [
        "Executive Summary",
        "Business Context and Drivers",
        "Architecture Scope",
        "Stakeholder Summary",
        "Baseline Architecture Overview",
        "Target Architecture Overview",
        "Gap Summary",
        "Key Risks and Assumptions",
        "Approval and Sign-off",
    ],
    "gap-analysis": [
        "Introduction",
        "Scope",
        "Baseline Architecture Summary",
        "Target Architecture Summary",
        "Gap Analysis",
        "Recommendations",
    ],
    "app-portfolio": [
        "Introduction",
        "Application Inventory",
        "Lifecycle Summary",
        "Recommendations",
    ],
    "requirements-register": [
        "Introduction",
        "Requirements Summary",
        "Requirements Register",
        "Traceability Matrix",
    ],
    "roadmap": [
        "Introduction",
        "Migration Strategy",
        "Transition Architectures",
        "Implementation Roadmap",
        "Dependencies and Risks",
    ],
    "stakeholder-map": [
        "Introduction",
        "Stakeholder Register",
        "Power/Interest Analysis",
        "Engagement Plan",
    ],
    "risk-register": [
        "Risk Summary",
        "Critical Risks",
        "High Risks",
        "Medium Risks",
        "Low Risks",
        "Closed / Accepted Risks",
        "Risk Heatmap Summary",
        "Source Artifact Cross-Reference",
    ],
}


# ---------------------------------------------------------------------------
# Document properties
# ---------------------------------------------------------------------------

def set_core_properties(doc, meta: dict, engagement: dict) -> None:
    """Write Dublin Core / Office core properties from artifact metadata."""
    props = doc.core_properties

    props.title           = meta.get("artifact") or engagement.get("name", "Architecture Document")
    props.subject         = f"ADM Phase {meta['phase']}" if meta.get("phase") else ""
    props.author          = engagement.get("sponsor") or engagement.get("author", "")
    props.last_modified_by = engagement.get("sponsor", "")
    props.category        = "Enterprise Architecture"
    props.content_status  = meta.get("status", "Draft")
    props.identifier      = meta.get("artifactId", "")
    props.description     = engagement.get("scope", "")

    # version field (string)
    props.version = str(meta.get("version", "0.1"))

    # keywords: TOGAF + phase + engagement type
    kw_parts = ["TOGAF", "Enterprise Architecture"]
    if meta.get("phase"):
        kw_parts.append(f"Phase {meta['phase']}")
    if engagement.get("engagementType"):
        kw_parts.append(engagement["engagementType"])
    props.keywords = ", ".join(kw_parts)

    # comments: engagement name + review status — visible in Windows Explorer
    comments_parts = [engagement.get("name", "")]
    if meta.get("reviewStatus"):
        comments_parts.append(f"Review: {meta['reviewStatus']}")
    props.comments = " | ".join(p for p in comments_parts if p)


def set_custom_properties(doc, props: dict) -> None:
    """
    Write EA-specific custom document properties (visible in Word File → Properties →
    Custom tab and searchable via Windows Explorer advanced search).

    Requires lxml (installed as a python-docx dependency). Skips silently if unavailable.
    """
    if not _HAS_LXML:
        return
    if not props:
        return

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    CUSTOM_NS   = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
    VT_NS       = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
    REL_TYPE    = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties"
    CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.custom-properties+xml"
    FMTID       = "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}"

    root = _etree.Element(
        f"{{{CUSTOM_NS}}}Properties",
        nsmap={None: CUSTOM_NS, "vt": VT_NS},
    )

    for pid, (name, value) in enumerate(props.items(), start=2):
        prop = _etree.SubElement(
            root,
            f"{{{CUSTOM_NS}}}property",
            fmtid=FMTID,
            pid=str(pid),
            name=name,
        )
        vt_str = _etree.SubElement(prop, f"{{{VT_NS}}}lpwstr")
        vt_str.text = str(value) if value is not None else ""

    xml_bytes = _etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    custom_part = Part(
        PackURI("/docProps/custom.xml"),
        CONTENT_TYPE,
        xml_bytes,
        doc.part.package,
    )
    doc.part.package.relate_to(custom_part, REL_TYPE)


# ---------------------------------------------------------------------------
# Document layout helpers
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_cover_page(doc, title, org, architect, doc_type, version_label, doc_date):
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(org)
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_type)
    run.font.size = Pt(14)
    run.font.color.rgb = MED_BLUE

    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    meta_rows = [
        ("Prepared by", architect),
        ("Organisation", org),
        ("Date", doc_date),
        ("Version", version_label),
    ]
    for i, (label, value) in enumerate(meta_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value

    doc.add_page_break()


def add_toc_placeholder(doc):
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = DARK_BLUE
    p = doc.add_paragraph("[Update table of contents after opening in Word: References → Update Table]")
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.runs[0].font.italic = True
    doc.add_page_break()


def add_section(doc, heading, content="", level=1):
    h = doc.add_heading(heading, level=level)
    color = DARK_BLUE if level == 1 else MED_BLUE
    for run in h.runs:
        run.font.color.rgb = color
    if content:
        doc.add_paragraph(content)
    else:
        p = doc.add_paragraph("[To be completed]")
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.runs[0].font.italic = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_background(hdr_cells[i], "1F3864")

    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = str(cell_text)
        if ri % 2 == 0:
            for cell in row_cells:
                set_cell_background(cell, "D6E4F0")

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def build_document(args, content: dict, meta: dict, engagement: dict) -> None:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Cover page values from metadata ---
    title      = meta.get("artifact") or args.title or engagement.get("name", "Architecture Document")
    org        = args.org or engagement.get("organisation", "Organisation")
    architect  = args.architect or engagement.get("sponsor", "Architect")
    status     = meta.get("status", "Draft")
    version    = str(meta.get("version", "0.1"))
    version_label = f"{version} {status}"

    last_modified = meta.get("lastModified") or ""
    try:
        doc_date = datetime.fromisoformat(last_modified).strftime("%B %Y") if last_modified else date.today().strftime("%B %Y")
    except ValueError:
        doc_date = date.today().strftime("%B %Y")

    phase    = meta.get("phase", "")
    doc_type = f"ADM Phase {phase} — Enterprise Architecture" if phase else "Enterprise Architecture"

    add_cover_page(doc, title, org, architect, doc_type, version_label, doc_date)
    add_toc_placeholder(doc)

    # --- Content sections ---
    artifact_type = args.type
    sections = content.get("sections", [])

    if sections:
        for sec in sections:
            heading = sec.get("heading", "Section")
            body    = sec.get("content", "")
            level   = sec.get("level", 1)
            add_section(doc, heading, body, level)
            if "table" in sec:
                t = sec["table"]
                add_table(doc, t.get("headers", []), t.get("rows", []))
    else:
        default_sections = ARTIFACT_SECTIONS.get(artifact_type, ["Content"])
        for sec_name in default_sections:
            add_section(doc, sec_name)

    for tbl in content.get("tables", []):
        heading = tbl.get("heading", "")
        if heading:
            add_section(doc, heading, level=2)
        add_table(doc, tbl.get("headers", []), tbl.get("rows", []))

    # --- Core properties (Dublin Core / Office standard) ---
    set_core_properties(doc, meta, engagement)

    # --- Custom properties (EA-specific, visible in Word Properties → Custom) ---
    custom_props = {}
    if meta.get("artifactId"):
        custom_props["EA.ArtifactId"]      = meta["artifactId"]
    if meta.get("artifact"):
        custom_props["EA.ArtifactType"]    = meta["artifact"]
    if meta.get("phase"):
        custom_props["EA.Phase"]           = meta["phase"]
    if meta.get("status"):
        custom_props["EA.Status"]          = meta["status"]
    if meta.get("reviewStatus"):
        custom_props["EA.ReviewStatus"]    = meta["reviewStatus"]
    if engagement.get("slug"):
        custom_props["EA.EngagementSlug"]  = engagement["slug"]
    if engagement.get("name"):
        custom_props["EA.EngagementName"]  = engagement["name"]
    if engagement.get("engagementType"):
        custom_props["EA.EngagementType"]  = engagement["engagementType"]

    set_custom_properties(doc, custom_props)

    doc.save(args.output)
    print(f"✅ Word document saved: {args.output}")
    if custom_props:
        print(f"   Core properties: title, subject, author, keywords, content_status, version, identifier")
        print(f"   Custom properties: {', '.join(custom_props.keys())}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate EA Assistant Word document")
    parser.add_argument("--type",           required=True,  help="Artifact type (vision, gap-analysis, etc.)")
    parser.add_argument("--title",          default=None,   help="Document title")
    parser.add_argument("--org",            default=None,   help="Organisation name")
    parser.add_argument("--architect",      default=None,   help="Lead architect name")
    parser.add_argument("--output",         required=True,  help="Output filename (.docx)")
    parser.add_argument("--content",        default="{}",   help="JSON content string or @file.json")
    parser.add_argument("--engagement-dir", default=None,   help="Path to engagement directory containing engagement.json")
    args = parser.parse_args()

    # Load engagement metadata
    engagement = {}
    if args.engagement_dir:
        engagement_path = os.path.join(args.engagement_dir, "engagement.json")
        if os.path.exists(engagement_path):
            with open(engagement_path) as f:
                engagement = json.load(f)
            if not args.title:
                args.title = engagement.get("name", "Architecture Document")
            if not args.org:
                args.org = engagement.get("organisation", "Organisation")
            if not args.architect:
                args.architect = engagement.get("sponsor", "Architect")
        else:
            print(f"WARNING: engagement.json not found at {engagement_path}")

    if not args.title:
        print("ERROR: --title is required when --engagement-dir is not provided")
        sys.exit(1)
    if not args.org:
        args.org = "Organisation"
    if not args.architect:
        args.architect = "Architect"

    # Load content JSON
    content_str = args.content
    if content_str.startswith("@"):
        with open(content_str[1:]) as f:
            content_str = f.read()
    try:
        content = json.loads(content_str)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON content — {e}")
        sys.exit(1)

    # Extract artifact metadata from content JSON (populated by ea-generate.md Step 3)
    meta = content.pop("meta", {})

    build_document(args, content, meta, engagement)


if __name__ == "__main__":
    main()
