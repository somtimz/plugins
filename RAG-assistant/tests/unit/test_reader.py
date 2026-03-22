"""Tests for scripts/lib/reader.py"""
import sys
import io
from pathlib import Path
from unittest.mock import MagicMock, patch, mock_open, PropertyMock

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent.parent / "scripts"))

from lib.reader import read_document, UnreadableError, EmptyDocumentError, FileSizeLimitError


# ---------------------------------------------------------------------------
# Plain-text formats
# ---------------------------------------------------------------------------

class TestTxtFiles:
    def test_txt_returns_full_utf8_text(self, tmp_path):
        content = "Hello, world!\nThis is a test.\nUnicode: café résumé"
        f = tmp_path / "doc.txt"
        f.write_text(content, encoding="utf-8")
        result = read_document(str(f), max_file_size_mb=10)
        assert result == content

    def test_md_returns_full_utf8_text(self, tmp_path):
        content = "# Heading\n\nSome **bold** text and `code`."
        f = tmp_path / "doc.md"
        f.write_text(content, encoding="utf-8")
        result = read_document(str(f), max_file_size_mb=10)
        assert result == content

    def test_empty_txt_raises_empty_document_error(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        with pytest.raises(EmptyDocumentError):
            read_document(str(f), max_file_size_mb=10)


# ---------------------------------------------------------------------------
# PDF files
# ---------------------------------------------------------------------------

class TestPdfFiles:
    def test_pdf_with_text_returns_extracted_content(self, tmp_path):
        fake_page = MagicMock()
        fake_page.extract_text.return_value = "Page one text."
        fake_reader = MagicMock()
        fake_reader.is_encrypted = False
        fake_reader.pages = [fake_page]

        f = tmp_path / "doc.pdf"
        f.write_bytes(b"%PDF-1.4 fake content")

        with patch("lib.reader.PdfReader", return_value=fake_reader):
            result = read_document(str(f), max_file_size_mb=10)

        assert "Page one text." in result

    def test_encrypted_pdf_raises_unreadable_error(self, tmp_path):
        import pypdf.errors

        f = tmp_path / "encrypted.pdf"
        f.write_bytes(b"%PDF-1.4 encrypted fake")

        with patch("lib.reader.PdfReader", side_effect=pypdf.errors.PdfReadError("encrypted")):
            with pytest.raises(UnreadableError):
                read_document(str(f), max_file_size_mb=10)

    def test_scanned_pdf_no_text_raises_empty_document_error(self, tmp_path):
        fake_page = MagicMock()
        fake_page.extract_text.return_value = ""
        fake_reader = MagicMock()
        fake_reader.is_encrypted = False
        fake_reader.pages = [fake_page, fake_page]

        f = tmp_path / "scanned.pdf"
        f.write_bytes(b"%PDF-1.4 scanned fake")

        with patch("lib.reader.PdfReader", return_value=fake_reader):
            with pytest.raises(EmptyDocumentError):
                read_document(str(f), max_file_size_mb=10)


# ---------------------------------------------------------------------------
# DOCX files
# ---------------------------------------------------------------------------

class TestDocxFiles:
    def test_docx_returns_joined_paragraph_text(self, tmp_path):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument()
            doc.add_paragraph("First paragraph.")
            doc.add_paragraph("Second paragraph.")
            docx_path = tmp_path / "doc.docx"
            doc.save(str(docx_path))
            result = read_document(str(docx_path), max_file_size_mb=10)
            assert "First paragraph." in result
            assert "Second paragraph." in result
        except ImportError:
            # Fall back to mocking if python-docx not installed
            fake_para1 = MagicMock()
            fake_para1.text = "First paragraph."
            fake_para2 = MagicMock()
            fake_para2.text = "Second paragraph."
            fake_doc = MagicMock()
            fake_doc.paragraphs = [fake_para1, fake_para2]

            f = tmp_path / "doc.docx"
            f.write_bytes(b"PK fake docx")

            with patch("lib.reader.Document", return_value=fake_doc):
                result = read_document(str(f), max_file_size_mb=10)

            assert "First paragraph." in result
            assert "Second paragraph." in result

    def test_docx_via_mock(self, tmp_path):
        fake_para1 = MagicMock()
        fake_para1.text = "Alpha text."
        fake_para2 = MagicMock()
        fake_para2.text = "Beta text."
        fake_doc = MagicMock()
        fake_doc.paragraphs = [fake_para1, fake_para2]

        f = tmp_path / "mock.docx"
        f.write_bytes(b"PK fake docx content")

        with patch("lib.reader.Document", return_value=fake_doc):
            result = read_document(str(f), max_file_size_mb=10)

        assert "Alpha text." in result
        assert "Beta text." in result


# ---------------------------------------------------------------------------
# Unsupported / wrong formats
# ---------------------------------------------------------------------------

class TestUnsupportedFormats:
    def test_doc_file_raises_unreadable_error_with_docx_message(self, tmp_path):
        f = tmp_path / "legacy.doc"
        f.write_bytes(b"\xd0\xcf\x11\xe0 fake OLE content")
        with pytest.raises(UnreadableError, match="docx"):
            read_document(str(f), max_file_size_mb=10)

    def test_xlsx_raises_unreadable_error(self, tmp_path):
        f = tmp_path / "spreadsheet.xlsx"
        f.write_bytes(b"PK fake xlsx content")
        with pytest.raises(UnreadableError):
            read_document(str(f), max_file_size_mb=10)


# ---------------------------------------------------------------------------
# File-not-found
# ---------------------------------------------------------------------------

class TestFileNotFound:
    def test_missing_file_raises_file_not_found_error(self, tmp_path):
        missing = str(tmp_path / "does_not_exist.txt")
        with pytest.raises(FileNotFoundError):
            read_document(missing, max_file_size_mb=10)


# ---------------------------------------------------------------------------
# File size limit
# ---------------------------------------------------------------------------

class TestFileSizeLimit:
    def test_oversized_file_raises_file_size_limit_error(self, tmp_path):
        f = tmp_path / "big.txt"
        # Write 1 byte — limit will be set to effectively 0 MB (near zero)
        f.write_text("x", encoding="utf-8")
        # max_file_size_mb=0 is effectively below any real file size,
        # but implementations may guard against 0; use a very small limit
        # by monkeypatching os.path.getsize to return a huge value.
        with patch("os.path.getsize", return_value=200 * 1024 * 1024):
            with pytest.raises(FileSizeLimitError):
                read_document(str(f), max_file_size_mb=10)

    def test_oversized_file_never_opens_file(self, tmp_path):
        f = tmp_path / "big.txt"
        f.write_text("some content", encoding="utf-8")

        m = mock_open()
        with patch("os.path.getsize", return_value=200 * 1024 * 1024):
            with patch("builtins.open", m):
                with pytest.raises(FileSizeLimitError):
                    read_document(str(f), max_file_size_mb=10)

        m.assert_not_called()
